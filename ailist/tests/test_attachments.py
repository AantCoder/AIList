"""
Unit tests for AIListBase._file_to_langchain, compile_block, and FileConverter.

WHAT IS TESTED:
  T1-T12: Every file type (text, html, office, pdf, image, audio, video) passes through
    _file_to_langchain and returns blocks with the correct structure for each provider.
    compile_block sets is_inline correctly. can_accept reflects provider capabilities.
    URL blocks, text_file block, unknown extension error.
  T13:  describe_image -- LLaVA via Ollama (mocked + live).
  T14:  FileConverter.pdf_to_text      -- basic extraction (pypdf).
  T15:  FileConverter.office_to_text   -- docx, xlsx, pptx; sheet/slide labels; unknown ext.
  T16:  FileConverter.html_to_text     -- tags stripped (BeautifulSoup).
  T17:  FileConverter.audio_to_transcript -- mp3, wav; language param (faster-whisper).
  T18:  FileConverter.export_pdf       -- plain text, markdown, explicit path (xhtml2pdf).
  T19:  FileConverter.export_docx      -- headings parsed (#/##/###), title, explicit path.
  T20:  FileConverter.export_xlsx      -- list-of-lists, list-of-dicts, sheet name, explicit path.

WHAT IS NOT TESTED:
  - Real LLM calls (no network requests).
  - Audio/video transcription quality -- only that output is non-empty.
  - Video fallback (frames + transcription) -- requires opencv and moviepy.
  - QWEN and TEXT_ONLY_WITH_SYSTEM providers -- caps covered by GEMINI and TEXT_ONLY.
  - compile_section / compile_combine -- integration level above this scope.
  - .doc (legacy format) -- requires LibreOffice installed.
  - Real LLaVA calls (Ollama mocked in T13; live smoke test optional).
  - export_pdf with complex HTML input -- covered by plain text and markdown subtests.
  - FileConverter.export_xlsx formula/chart support -- not implemented (use fastskills xlsx).

REQUIRED FILES IN TEST_FILES_DIR:
  sample.txt, sample.py, sample.json, sample.csv, sample.xml, sample.md
  sample.html, sample.pdf, sample.docx, sample.xlsx, sample.pptx
  sample.jpg, sample.png, sample.mp3, sample.wav, sample.mp4
"""

import sys
import os
import importlib.util
import time
from pathlib import Path

# Path to test fixture files.
# Set AILIST_TEST_FILES_DIR env var, or pass as first argument: python test_attachments.py /path/to/files
TEST_FILES_DIR = (
    sys.argv[1] if len(sys.argv) > 1 else
    os.environ.get("AILIST_TEST_FILES_DIR", "test_files")
)

from ailist import AIListBase, AIList, Provider, ProviderCaps, EXT_TYPE, CompiledBlock

# --------------------------------------------------------------
# Helpers
# --------------------------------------------------------------

def fp(name: str) -> str:
    """Return full path to a test fixture file."""
    return str(Path(TEST_FILES_DIR) / name)

def make_ai(provider: ProviderCaps) -> AIListBase:
    """Create an AIListBase with the given provider without initializing the agent."""
    ai = AIList.__new__(AIList)
    AIListBase.__init__(ai, "dummy", 99999, provider)
    # Примечание: transcript_on_cuda здесь не выставляем.
    # _converter создаётся внутри AIListBase.__init__ до того, как мы можем изменить флаг,
    # поэтому установка ai.transcript_on_cuda=True здесь не влияет на use_cuda в _converter.
    # Для тестов транскрипции используется CPU (дефолт), что достаточно для проверки корректности.
    return ai

_any_fail = False
_t_total  = time.perf_counter()   # общий таймер всего прогона
_t_section = None                  # таймер текущей секции

def check(cond: bool, msg: str):
    global _any_fail
    print(f"  {'OK  ' if cond else 'FAIL'} | {msg}")
    if not cond:
        _any_fail = True

def section(title: str):
    global _t_section
    # Печатаем время предыдущей секции если она была
    if _t_section is not None:
        print(f"  Time: {time.perf_counter() - _t_section:.3f} sec")
    _t_section = time.perf_counter()
    print(f"\n{'='*60}")
    print(f"  {title}")
    print("=" * 60)

def sub(title: str):
    print(f"\n  [{title}]")


# --------------------------------------------------------------
# T1. Plain text files: txt, py, json, csv, xml, md
# --------------------------------------------------------------

section("T1. Plain text files")
print("  Expected: [label, content] blocks, both type=text, content non-empty, is_inline=False")

ai = make_ai(Provider.ANTHROPIC)

for fname in ["sample.txt", "sample.py", "sample.json", "sample.csv", "sample.xml", "sample.md"]:
    path = fp(fname)
    ext = fname.rsplit(".", 1)[-1]
    sub(fname)
    try:
        blocks = ai._file_to_langchain(path, ext)
        cb = ai.compile_block({"file": path})
        check(len(blocks) == 2,                    "exactly two blocks: [label, content]")
        check(blocks[0]["type"] == "text",          "label block: type=text")
        check(fname in blocks[0]["text"],           "label block: filename present")
        check(blocks[1]["type"] == "text",          "content block: type=text")
        check(bool(blocks[1]["text"].strip()),      "content block: non-empty")
        check(cb.is_inline is False,                "is_inline=False (named attachment)")
    except Exception as e:
        print(f"  EXCEPTION: {e}")


# --------------------------------------------------------------
# T2. HTML -> text via BeautifulSoup
# --------------------------------------------------------------

section("T2. HTML -> text via BeautifulSoup")
print("  Expected: is_inline=True, type=text, no HTML tags in output")

sub("sample.html / TEXT_ONLY")
ai = make_ai(Provider.TEXT_ONLY)
try:
    blocks = ai._file_to_langchain(fp("sample.html"), "html")
    label_text = blocks[0]["text"]
    content_text = blocks[1]["text"]
    check(len(blocks) == 2,                        "exactly two blocks: [label, content]")
    check(blocks[0]["type"] == "text",              "label block: type=text")
    check("sample.html" in label_text,             "label block: filename present")
    check(blocks[1]["type"] == "text",              "content block: type=text")
    check("<html" not in content_text.lower(),      "HTML tags stripped")
    check(bool(content_text.strip()),               "content is non-empty")
    cb = ai.compile_block({"file": fp("sample.html")})
    check(cb.is_inline is False,                    "is_inline=False (named attachment)")
    print(f"    Text preview: {content_text[:120].strip()!r}")
except Exception as e:
    print(f"  EXCEPTION: {e}")

# HTML conversion is provider-independent - verify ANTHROPIC gives the same result
sub("sample.html / ANTHROPIC")
try:
    ai2 = make_ai(Provider.ANTHROPIC)
    blocks2 = ai2._file_to_langchain(fp("sample.html"), "html")
    check(blocks2[0]["type"] == "text", "ANTHROPIC also returns text (not binary)")
except Exception as e:
    print(f"  EXCEPTION: {e}")


# --------------------------------------------------------------
# T3. Office documents -> text: docx, xlsx, pptx
# --------------------------------------------------------------

section("T3. Office -> text (docx, xlsx, pptx)")
print("  Expected: is_inline=True, type=text, non-empty, conversion is provider-independent")

for fname, ext in [("sample.docx", "docx"), ("sample.xlsx", "xlsx"), ("sample.pptx", "pptx")]:
    sub(fname)
    for pname, prov in [("ANTHROPIC", Provider.ANTHROPIC), ("TEXT_ONLY", Provider.TEXT_ONLY)]:
        ai = make_ai(prov)
        try:
            blocks = ai._file_to_langchain(fp(fname), ext)
            check(len(blocks) == 2,                    f"{pname}: exactly two blocks: [label, content]")
            check(blocks[0]["type"] == "text",          f"{pname}: label block: type=text")
            check(fname in blocks[0]["text"],           f"{pname}: label block: filename present")
            check(blocks[1]["type"] == "text",          f"{pname}: content block: type=text")
            check(bool(blocks[1]["text"].strip()),      f"{pname}: content block: non-empty")
            cb = ai.compile_block({"file": fp(fname)})
            check(cb.is_inline is False,                f"{pname}: is_inline=False (named attachment)")
            print(f"    {pname} preview: {blocks[1]['text'][:80].strip()!r}")
        except Exception as e:
            print(f"    {pname} EXCEPTION: {e}")


# --------------------------------------------------------------
# T4. PDF
# --------------------------------------------------------------

section("T4. PDF")
print("  Expected: Anthropic->file block, Gemini->document base64,")
print("           OPENAI/QWEN/TEXT_ONLY->text via pypdf (transparent conversion, no stub)")

sub("ANTHROPIC - pdf_format='file'")
ai = make_ai(Provider.ANTHROPIC)
try:
    blocks = ai._file_to_langchain(fp("sample.pdf"), "pdf")
    # blocks = [label, file_block]
    label_b = blocks[0]
    b = blocks[1]
    check(len(blocks) == 2,                                                  "exactly two blocks: [label, file]")
    check(label_b["type"] == "text",                                         "label block: type=text")
    check("sample.pdf" in label_b["text"],                                   "label block: filename present")
    check(b["type"] == "file",                                               "type=file")
    check("filename" in b.get("file", {}),                                   "file.filename present")
    _fdata = b["file"].get("file_data", "")
    check(_fdata.startswith("data:application/pdf;base64,"),
          f"file_data is valid (got prefix: {_fdata[:40]!r})")
    cb = ai.compile_block({"file": fp("sample.pdf")})
    check(cb.is_inline is False,                                             "is_inline=False (binary attachment)")
except Exception as e:
    print(f"  EXCEPTION: {e}")

sub("GEMINI - pdf_format='document'")
ai = make_ai(Provider.GEMINI)
try:
    blocks = ai._file_to_langchain(fp("sample.pdf"), "pdf")
    # blocks = [label, document_block]
    b = blocks[1]
    check(len(blocks) == 2,                            "exactly two blocks: [label, document]")
    check(b["type"] == "document",                     "type=document")
    check(b.get("source", {}).get("type") == "base64", "source.type=base64")
    check(bool(b["source"].get("data")),               "data is non-empty")
except Exception as e:
    print(f"  EXCEPTION: {e}")

# pdf_format=None (OPENAI, QWEN, TEXT_ONLY): transparent pypdf conversion for all
sub("OPENAI / QWEN / TEXT_ONLY - pdf_format=None -> text via pypdf (transparent)")
for pname, prov in [("OPENAI", Provider.OPENAI), ("QWEN", Provider.QWEN), ("TEXT_ONLY", Provider.TEXT_ONLY)]:
    ai = make_ai(prov)
    try:
        blocks = ai._file_to_langchain(fp("sample.pdf"), "pdf")
        # blocks = [pdf_label, content]  (FILE_PDF_TEXT_LABEL used as label)
        label_b = blocks[0]
        content_b = blocks[1]
        check(len(blocks) == 2,                  f"{pname}: exactly two blocks: [label, content]")
        check(label_b["type"] == "text",          f"{pname}: label block: type=text")
        check("PDF" in label_b["text"],           f"{pname}: FILE_PDF_TEXT_LABEL present in label")
        check(content_b["type"] == "text",        f"{pname}: content block: type=text")
        check(bool(content_b["text"].strip()),    f"{pname}: content block: non-empty")
        print(f"    {pname} preview: {label_b['text'][:100].strip()!r}")
    except Exception as e:
        print(f"    {pname} EXCEPTION: {e}")


# --------------------------------------------------------------
# T5. Images: jpg, png
# --------------------------------------------------------------

section("T5. Images (jpg, png)")
print("  Expected: ANTHROPIC/OPENAI->image_url block, GEMINI->media block, TEXT_ONLY->stub text")

for fname, ext in [("sample.jpg", "jpg"), ("sample.png", "png")]:

    sub(f"{fname} / ANTHROPIC - image_url")
    ai = make_ai(Provider.ANTHROPIC)
    try:
        blocks = ai._file_to_langchain(fp(fname), ext)
        # blocks = [label, image_url_block]
        b = blocks[1]
        mime = "image/jpeg" if ext == "jpg" else "image/png"
        check(len(blocks) == 2,                                         "exactly two blocks: [label, image]")
        check(blocks[0]["type"] == "text",                              "label block: type=text")
        check(fname in blocks[0]["text"],                               "label block: filename present")
        check(b["type"] == "image_url",                                 "type=image_url")
        _iurl = b["image_url"].get("url", "")
        check(_iurl.startswith(f"data:{mime};base64,"),
              f"url is valid data URI (got prefix: {_iurl[:50]!r})")
        cb = ai.compile_block({"file": fp(fname)})
        check(cb.is_inline is False,                                    "is_inline=False")
    except Exception as e:
        print(f"    EXCEPTION: {e}")

    sub(f"{fname} / GEMINI - media")
    ai = make_ai(Provider.GEMINI)
    try:
        blocks = ai._file_to_langchain(fp(fname), ext)
        # blocks = [label, media_block]
        b = blocks[1]
        check(len(blocks) == 2,        "exactly two blocks: [label, media]")
        check(b["type"] == "media",    "type=media")
        check("mime_type" in b,        "mime_type present")
        check(bool(b.get("data")),     "data is non-empty")
    except Exception as e:
        print(f"    EXCEPTION: {e}")

    sub(f"{fname} / TEXT_ONLY - LLaVA description (or stub if Ollama unavailable)")
    ai = make_ai(Provider.TEXT_ONLY)
    try:
        blocks = ai._file_to_langchain(fp(fname), ext)
        b = blocks[0]
        check(b["type"] == "text", "type=text")
        check(bool(b["text"]),     "text is non-empty")
        # Два возможных исхода: LLaVA вернула описание или Ollama недоступна -> заглушка.
        # В обоих случаях блок должен быть текстовым.
        llava_path   = len(blocks) == 2 and "LLaVA" in blocks[0].get("text", "")
        stub_path    = len(blocks) == 1
        check(llava_path or stub_path, "either LLaVA description (2 blocks) or stub (1 block)")
        if llava_path:
            print(f"    LLaVA path: description present, preview: {blocks[1]['text'][:80].strip()!r}")
        else:
            print(f"    Stub path (Ollama unavailable): {b['text']!r}")
    except Exception as e:
        print(f"    EXCEPTION: {e}")


# --------------------------------------------------------------
# T6. Audio: mp3, wav
# --------------------------------------------------------------

section("T6. Audio (mp3, wav)")
print("  Expected: Anthropic->Whisper text (audio_format=None), Gemini->media, OpenAI->input_audio,")
print("           QWEN/TEXT_ONLY->Whisper transcription")

for fname, ext in [("sample.mp3", "mp3"), ("sample.wav", "wav")]:

    sub(f"{fname} / ANTHROPIC - Whisper fallback (audio_format=None)")
    ai = make_ai(Provider.ANTHROPIC)
    try:
        blocks = ai._file_to_langchain(fp(fname), ext)
        # blocks = [audio_label, transcript_text]
        check(len(blocks) == 2,                    "exactly two blocks: [label, transcript]")
        check(blocks[0]["type"] == "text",          "label block: type=text")
        check("audio transcription" in blocks[0]["text"], "label: FILE_AUDIO_TRANSCRIPT marker")
        check(blocks[1]["type"] == "text",          "transcript block: type=text")
        check(bool(blocks[1]["text"].strip()),      "transcript: non-empty")
    except Exception as e:
        print(f"    EXCEPTION: {e}")

    sub(f"{fname} / GEMINI - media")
    ai = make_ai(Provider.GEMINI)
    try:
        blocks = ai._file_to_langchain(fp(fname), ext)
        # blocks = [label, media_block]
        b = blocks[1]
        check(len(blocks) == 2,                      "exactly two blocks: [label, media]")
        check(b["type"] == "media",                  "type=media")
        check(bool(b.get("data")),                   "data is non-empty")
        check("audio" in b.get("mime_type", ""),     "mime_type contains 'audio'")
    except Exception as e:
        print(f"    EXCEPTION: {e}")

    sub(f"{fname} / OPENAI - input_audio")
    ai = make_ai(Provider.OPENAI)
    try:
        blocks = ai._file_to_langchain(fp(fname), ext)
        # blocks = [label, input_audio_block]
        b = blocks[1]
        check(len(blocks) == 2,                              "exactly two blocks: [label, audio]")
        check(b["type"] == "input_audio",                    "type=input_audio")
        check(b["input_audio"]["format"] in ("mp3", "wav"),  "format is mp3 or wav")
        check(bool(b["input_audio"].get("data")),            "data is non-empty")
    except Exception as e:
        print(f"    EXCEPTION: {e}")

    sub(f"{fname} / TEXT_ONLY - Whisper fallback")
    ai = make_ai(Provider.TEXT_ONLY)
    try:
        blocks = ai._file_to_langchain(fp(fname), ext)
        # blocks = [audio_label, transcript_text]
        check(len(blocks) == 2,                    "exactly two blocks: [label, transcript]")
        check(blocks[0]["type"] == "text",          "label block: type=text")
        check("audio transcription" in blocks[0]["text"], "label: FILE_AUDIO_TRANSCRIPT marker")
        check(blocks[1]["type"] == "text",          "transcript block: type=text")
        check(bool(blocks[1]["text"].strip()),      "transcript: non-empty")
        print(f"    Transcript label: {blocks[0]['text'][:120].strip()!r}")
    except Exception as e:
        print(f"    EXCEPTION: {e}")


# --------------------------------------------------------------
# T7. Video - native support: ANTHROPIC, QWEN
# --------------------------------------------------------------

section("T7. Video / native support (QWEN)")
print("  Expected: QWEN->video_url block")
print("  Note: ANTHROPIC has supports_video=False - uses frame fallback (tested in T8)")

# ANTHROPIC does not support native video (supports_video=False) - fallback to frames.
# Native video is only available for QWEN (video_url format).
sub("sample.mp4 / QWEN - video_url")
ai = make_ai(Provider.QWEN)
try:
    blocks = ai._file_to_langchain(fp("sample.mp4"), "mp4")
    # blocks = [label, video_url_block]
    b = blocks[1]
    check(len(blocks) == 2,                                                  "exactly two blocks: [label, video]")
    check(blocks[0]["type"] == "text",                                       "label block: type=text")
    check("sample.mp4" in blocks[0]["text"],                                 "label block: filename present")
    check(b["type"] == "video_url",                                          "type=video_url")
    check(b["video_url"]["url"].startswith("data:video/mp4;base64,"),        "url carries video MIME type")
except Exception as e:
    print(f"  EXCEPTION: {e}")


# --------------------------------------------------------------
# T8. Video - fallback (frames + transcription) for OPENAI
# --------------------------------------------------------------

section("T8. Video / fallback (OPENAI and ANTHROPIC - no native video support)")
print("  Expected: image_url frame blocks, optionally a text transcription block")
print("  Note: requires opencv-python; without moviepy/OPENAI_API_KEY only frames are extracted")

_cv2_available = importlib.util.find_spec("cv2") is not None

if not _cv2_available:
    print("  SKIP - cv2 not installed, frame extraction unavailable")
else:
    ai = make_ai(Provider.OPENAI)
    try:
        blocks = ai._file_to_langchain(fp("sample.mp4"), "mp4")
        image_blocks = [b for b in blocks if b.get("type") == "image_url"]
        text_blocks  = [b for b in blocks if b.get("type") == "text"]
        # Expected structure: [transcript_label, transcript_text, frames_label, frame1, frame2, ...]
        # Labels are text blocks; frames are image_url blocks
        check(len(image_blocks) > 0, f"frame blocks present: {len(image_blocks)} block(s)")
        check(len(blocks) > 0,       "total blocks > 0")
        for ib in image_blocks:
            check(ib["image_url"]["url"].startswith("data:image/jpeg;base64,"),
                  "frame is a valid base64 JPEG")
            break  # checking the first frame is sufficient
        label_texts = [b["text"] for b in text_blocks if b.get("type") == "text"]
        frames_label = any("frames from video" in t for t in label_texts)
        check(frames_label, "frames label (FILE_VIDEO_FRAMES_LABEL) present")
        if text_blocks:
            print(f"    Found {len(text_blocks)} text block(s) (labels + optional transcript)")
    except Exception as e:
        print(f"  EXCEPTION: {e}")


# --------------------------------------------------------------
# T9. URL blocks
# --------------------------------------------------------------

section("T9. URL blocks")
print("  Expected: correct block type built without downloading the file")

TEST_IMAGE_URL = "https://example.com/photo.jpg"
TEST_PDF_URL   = "https://example.com/document.pdf"

sub("URL image / ANTHROPIC -> image_url")
ai = make_ai(Provider.ANTHROPIC)
blocks = ai._url_to_langchain(TEST_IMAGE_URL, "jpg")
# blocks = [label, image_url_block]
check(len(blocks) == 2,                                "exactly two blocks: [label, image]")
check(blocks[1]["type"] == "image_url",                "type=image_url")
check(blocks[1]["image_url"]["url"] == TEST_IMAGE_URL, "url passed through unchanged")

sub("URL image / GEMINI -> media with file_uri")
ai = make_ai(Provider.GEMINI)
blocks = ai._url_to_langchain(TEST_IMAGE_URL, "jpg")
# blocks = [label, media_block]
check(len(blocks) == 2,                            "exactly two blocks: [label, media]")
check(blocks[1]["type"] == "media",                "type=media")
check(blocks[1].get("file_uri") == TEST_IMAGE_URL, "file_uri equals original URL")
check("image" in blocks[1].get("mime_type", ""),   "mime_type contains 'image'")

sub("URL pdf / ANTHROPIC -> file")
ai = make_ai(Provider.ANTHROPIC)
blocks = ai._url_to_langchain(TEST_PDF_URL, "pdf")
# blocks = [label, file_block]
check(len(blocks) == 2,            "exactly two blocks: [label, file]")
check(blocks[1]["type"] == "file", "type=file")

sub("URL pdf / GEMINI -> document url")
ai = make_ai(Provider.GEMINI)
blocks = ai._url_to_langchain(TEST_PDF_URL, "pdf")
# blocks = [label, document_block]
check(len(blocks) == 2,                            "exactly two blocks: [label, document]")
check(blocks[1]["type"] == "document",             "type=document")
check(blocks[1]["source"]["type"] == "url",        "source.type=url")
check(blocks[1]["source"]["url"] == TEST_PDF_URL,  "url is correct")

sub("URL image / TEXT_ONLY -> stub text")
ai = make_ai(Provider.TEXT_ONLY)
blocks = ai._url_to_langchain(TEST_IMAGE_URL, "jpg")
# no label for stubs - single text block with URL info
check(blocks[0]["type"] == "text", "type=text (stub)")
check("URL" in blocks[0]["text"],  "stub text contains 'URL'")


# --------------------------------------------------------------
# T10. text_file block
# --------------------------------------------------------------

section("T10. text_file block")
print("  Expected: is_inline=True, text contains the filename label")

ai = make_ai(Provider.ANTHROPIC)

sub("text_file - read from disk")
try:
    cb = ai.compile_block({"text_file": fp("sample.txt")})
    check(cb.is_inline is True,                  "is_inline=True")
    check(len(cb.blocks) == 1,                   "exactly one block")
    check(cb.blocks[0]["type"] == "text",        "type=text")
    check("sample.txt" in cb.blocks[0]["text"],  "filename label present")
    check(bool(cb.blocks[0]["text"].strip()),     "content is non-empty")
except Exception as e:
    print(f"  EXCEPTION: {e}")

sub("text_file - pre-converted content")
cb = ai.compile_block({"text_file": "report.xlsx", "content": "col1\tcol2\n1\t2"})
check(cb.is_inline is True,                  "is_inline=True")
check("report.xlsx" in cb.blocks[0]["text"], "filename label present")
check("col1" in cb.blocks[0]["text"],        "content passed through correctly")


# --------------------------------------------------------------
# T11. can_accept
# --------------------------------------------------------------

section("T11. can_accept - True means NATIVE (no conversion), False means converted transparently")
print("  can_accept returns True only when the file is sent to the model as-is (binary/native).")
print("  False means conversion happens but the file IS still processed - just not natively.")
print("  Only image without binary support has NO processing path at all.")

ai_ant = make_ai(Provider.ANTHROPIC)
ai_oai = make_ai(Provider.OPENAI)
ai_txt = make_ai(Provider.TEXT_ONLY)
ai_qwn = make_ai(Provider.QWEN)

sub("TEXT_ONLY - nothing is native (no binary support)")
check(not ai_txt.can_accept({"file": fp("sample.txt")}),   "txt: False (converted to text)")
check(not ai_txt.can_accept({"file": fp("sample.docx")}),  "docx: False (converted to text)")
check(not ai_txt.can_accept({"file": fp("sample.html")}),  "html: False (converted to text)")
check(not ai_txt.can_accept({"file": fp("sample.pdf")}),   "pdf: False (pypdf conversion)")
check(not ai_txt.can_accept({"file": fp("sample.mp3")}),   "mp3: False (Whisper conversion)")
check(not ai_txt.can_accept({"file": fp("sample.mp4")}),   "mp4: False (frames conversion)")
check(not ai_txt.can_accept({"file": fp("sample.jpg")}),   "jpg: False (no binary; LLaVA or stub)")
check(ai_txt.can_accept({"prompt": "hello"}),              "prompt: True (plain text, always native)")
check(ai_txt.can_accept({"text_file": fp("sample.txt")}),  "text_file: True (pre-processed)")

sub("ANTHROPIC - everything is native")
check(ai_ant.can_accept({"file": fp("sample.pdf")}),        "pdf: True (native file block)")
check(not ai_ant.can_accept({"file": fp("sample.mp3")}),    "mp3: False (audio_format=None, Whisper fallback)")
check(not ai_ant.can_accept({"file": fp("sample.mp4")}),    "mp4: False (supports_video=False, frame fallback)")
check(ai_ant.can_accept({"file": fp("sample.jpg")}),        "jpg: True (native image_url)")
check(not ai_ant.can_accept({"file": fp("sample.txt")}),    "txt: False (read as text, not binary)")
check(not ai_ant.can_accept({"file": fp("sample.docx")}),   "docx: False (converted to text)")

sub("OPENAI - native image+audio, fallback for pdf and video")
check(ai_oai.can_accept({"file": fp("sample.jpg")}),       "jpg: True (native image_url)")
check(ai_oai.can_accept({"file": fp("sample.mp3")}),       "mp3: True (native input_audio)")
check(not ai_oai.can_accept({"file": fp("sample.pdf")}),   "pdf: False (pypdf conversion)")
check(not ai_oai.can_accept({"file": fp("sample.mp4")}),   "mp4: False (frames conversion)")

sub("QWEN - native image+video, fallback for pdf and audio")
check(ai_qwn.can_accept({"file": fp("sample.jpg")}),       "jpg: True (native image_url)")
check(ai_qwn.can_accept({"file": fp("sample.mp4")}),       "mp4: True (native video_url)")
check(not ai_qwn.can_accept({"file": fp("sample.pdf")}),   "pdf: False (pypdf conversion)")
check(not ai_qwn.can_accept({"file": fp("sample.mp3")}),   "mp3: False (Whisper conversion)")


# --------------------------------------------------------------
# T12. Unknown file extension -> ValueError
# --------------------------------------------------------------

section("T12. Unknown extension -> ValueError")
print("  Expected: ValueError raised with the extension name in the message")

ai = make_ai(Provider.ANTHROPIC)
try:
    ai._file_to_langchain(fp("sample.xyz"), "xyz")
    print("  FAIL - no exception raised")
    _any_fail = True
except ValueError as e:
    check("xyz" in str(e), f"ValueError contains extension name: {e}")
except Exception as e:
    print(f"  FAIL - unexpected exception type: {e}")
    _any_fail = True


# --------------------------------------------------------------
# T13. describe_image -- LLaVA via Ollama (mocked + live)
# --------------------------------------------------------------

section("T13. describe_image -- LLaVA via Ollama")
print("  Mocked subtests: no real HTTP call, assert_called_once() confirms LLaVA was invoked.")
print("  Live subtest: requires Ollama running with llava model.")

import unittest.mock as _mock
import json as _json
import urllib.error as _urllib_err
import urllib.request as _urllib_req

def _make_mock_resp(response_text: str):
    mock_resp = _mock.MagicMock()
    mock_resp.read.return_value = _json.dumps({"response": response_text}).encode()
    mock_resp.__enter__ = lambda s: s
    mock_resp.__exit__ = _mock.MagicMock(return_value=False)
    return mock_resp

def _make_ollama_mock(response_text: str):
    """Patches urllib.request.urlopen; yields the mock so callers can assert_called_once()."""
    return _mock.patch("urllib.request.urlopen", return_value=_make_mock_resp(response_text))

def _fail_test(msg: str):
    """Регистрирует FAIL и печатает сообщение (не бросает исключение, чтобы продолжить прогон)."""
    global _any_fail
    _any_fail = True
    print(f"  FAIL | {msg}")

ai = make_ai(Provider.TEXT_ONLY)

# -- Mocked: response text ----------------------------------------------------
sub("describe_image - response text returned correctly [mocked]")
with _make_ollama_mock("A cat sitting on a red chair.") as mock_urlopen:
    result = ai._converter.describe_image(fp("sample.jpg"))
if not mock_urlopen.called:
    _fail_test("urlopen was never called -- LLaVA code path not reached")
else:
    check(mock_urlopen.call_count == 1,             "urlopen called exactly once")
    check(result == "A cat sitting on a red chair.", "response text returned correctly")

# -- Mocked: custom prompt in payload -----------------------------------------
sub("describe_image - custom prompt sent in payload [mocked]")
captured = {}
def _capturing_urlopen(req, **kw):
    captured["body"] = _json.loads(req.data.decode())
    return _make_mock_resp("ok")
with _mock.patch("urllib.request.urlopen", side_effect=_capturing_urlopen) as mock_urlopen:
    ai._converter.describe_image(fp("sample.jpg"), prompt="List all objects.")
if not mock_urlopen.called:
    _fail_test("urlopen was never called -- LLaVA code path not reached")
else:
    check(mock_urlopen.call_count == 1,                                       "urlopen called exactly once")
    check(captured.get("body", {}).get("prompt") == "List all objects.",      "custom prompt in payload")
    check("images" in captured.get("body", {}),                               "images key present")
    check(len(captured["body"]["images"]) == 1,                               "exactly one image")
    check(bool(captured["body"]["images"][0]),                                 "image data non-empty base64")

# -- Mocked: model override ----------------------------------------------------
sub("describe_image - model parameter override [mocked]")
captured_model = {}
def _capture_model(req, **kw):
    captured_model["model"] = _json.loads(req.data.decode()).get("model")
    return _make_mock_resp("ok")
with _mock.patch("urllib.request.urlopen", side_effect=_capture_model) as mock_urlopen:
    ai._converter.describe_image(fp("sample.jpg"), model="llava:13b")
if not mock_urlopen.called:
    _fail_test("urlopen was never called -- LLaVA code path not reached")
else:
    check(mock_urlopen.call_count == 1,                  "urlopen called exactly once")
    check(captured_model.get("model") == "llava:13b",    "custom model name sent in payload")

# -- Mocked: _file_to_langchain LLaVA path ------------------------------------
sub("_file_to_langchain TEXT_ONLY + mock Ollama - LLaVA path (2 blocks) [mocked]")
with _make_ollama_mock("A scenic mountain landscape.") as mock_urlopen:
    blocks = ai._file_to_langchain(fp("sample.jpg"), "jpg")
if not mock_urlopen.called:
    _fail_test(
        "urlopen was never called -- _file_to_langchain silently fell back to stub. "
        "LLaVA code path not reached (exception swallowed?)"
    )
else:
    check(mock_urlopen.call_count == 1,            "urlopen called exactly once")
    check(len(blocks) == 2,                        "exactly two blocks: [label, description]")
    check(blocks[0]["type"] == "text",             "label block: type=text")
    check("LLaVA" in blocks[0]["text"],            "label: FILE_IMAGE_DESCRIPTION marker")
    check("sample.jpg" in blocks[0]["text"],       "label: filename present")
    check(blocks[1]["type"] == "text",             "description block: type=text")
    check("mountain" in blocks[1]["text"],         "description block: content from LLaVA")

# -- Mocked: fallback on Ollama error -----------------------------------------
sub("_file_to_langchain TEXT_ONLY + Ollama error - fallback to stub (1 block) [mocked]")
with _mock.patch("urllib.request.urlopen", side_effect=_urllib_err.URLError("connection refused")):
    blocks = ai._file_to_langchain(fp("sample.jpg"), "jpg")
check(len(blocks) == 1,            "exactly one block (stub on Ollama failure)")
check(blocks[0]["type"] == "text", "stub block: type=text")
check(bool(blocks[0]["text"]),     "stub block: non-empty")
print(f"    Stub text: {blocks[0]['text']!r}")

# -- Live: real Ollama call ----------------------------------------------------
sub("describe_image - real Ollama call [live]")
_ollama_live_ok = False
try:
    _urllib_req.urlopen(
        _urllib_req.Request("http://localhost:11434", method="HEAD"),
        timeout=2,
    )
    _ollama_live_ok = True
except Exception:
    pass

if not _ollama_live_ok:
    _fail_test(
        "Ollama is not running -- live LLaVA test cannot be executed.\n"
        "  To fix, run in a separate terminal:\n"
        "    ollama serve\n"
        "    ollama pull llava\n"
        "  Then re-run the tests."
    )
else:
    try:
        description = ai._converter.describe_image(fp("sample.jpg"))
        check(bool(description.strip()), "description is non-empty")
        print(f"    LLaVA description of sample.jpg: {description[:200].strip()!r}")
    except Exception as e:
        _fail_test(f"describe_image raised an exception: {e}")

if _t_section is not None:
    print(f"  Time: {time.perf_counter() - _t_section:.3f} sec")
    _t_section = None  # сбрасываем: следующий section() не напечатает лишнее время

# ==============================================================
# FileConverter -- прямые unit-тесты (без LLM, без агента)
# ==============================================================
#
# Тесты T14-T20 проверяют методы FileConverter напрямую.
# Зависимости:
#   pip install pypdf python-docx openpyxl python-pptx beautifulsoup4 lxml xhtml2pdf
# Тесты пропускаются (SKIP) если нужная библиотека не установлена.
# ==============================================================

from ailist import FileConverter as _FileConverter
import importlib.util as _ilu
import tempfile as _tempfile

_fc = _FileConverter()   # один объект на все тесты

def _lib(name: str) -> bool:
    """True если пакет установлен."""
    return _ilu.find_spec(name) is not None


# --------------------------------------------------------------
# T14. FileConverter.pdf_to_text
# --------------------------------------------------------------

section("T14. FileConverter.pdf_to_text")
print("  Expected: non-empty text extracted from sample.pdf")
print("  Requires: pip install pypdf")

if not _lib("pypdf"):
    print("  SKIP -- pypdf not installed")
else:
    sub("pdf_to_text -- basic extraction")
    try:
        text = _fc.pdf_to_text(fp("sample.pdf"))
        check(isinstance(text, str),    "returns str")
        check(bool(text.strip()),       "non-empty")
        print(f"    Preview: {text[:120].strip()!r}")
    except Exception as e:
        print(f"  EXCEPTION: {e}")
        _any_fail = True


# --------------------------------------------------------------
# T15. FileConverter.office_to_text
# --------------------------------------------------------------

section("T15. FileConverter.office_to_text")
print("  Expected: non-empty text for docx, xlsx, pptx")
print("  Requires: pip install python-docx openpyxl python-pptx")

for fname, ext, lib in [
    ("sample.docx", "docx", "docx"),
    ("sample.xlsx", "xlsx", "openpyxl"),
    ("sample.pptx", "pptx", "pptx"),
]:
    sub(f"office_to_text -- {fname}")
    if not _lib(lib):
        print(f"    SKIP -- {lib} not installed")
        continue
    try:
        text = _fc.office_to_text(fp(fname), ext=ext)
        check(isinstance(text, str),    "returns str")
        check(bool(text.strip()),       "non-empty")
        print(f"    Preview: {text[:120].strip()!r}")
    except Exception as e:
        print(f"    EXCEPTION: {e}")
        _any_fail = True

sub("office_to_text -- xlsx: sheet label present")
if _lib("openpyxl"):
    try:
        text = _fc.office_to_text(fp("sample.xlsx"), ext="xlsx")
        check("[Sheet:" in text,    "sheet label [Sheet: ...] present in output")
    except Exception as e:
        print(f"    EXCEPTION: {e}")
        _any_fail = True
else:
    print("    SKIP -- openpyxl not installed")

sub("office_to_text -- pptx: slide label present")
if _lib("pptx"):
    try:
        text = _fc.office_to_text(fp("sample.pptx"), ext="pptx")
        check("[Slide " in text,    "slide label [Slide N] present in output")
    except Exception as e:
        print(f"    EXCEPTION: {e}")
        _any_fail = True
else:
    print("    SKIP -- python-pptx not installed")

sub("office_to_text -- unknown ext -> RuntimeError")
try:
    _fc.office_to_text(fp("sample.txt"), ext="xyz")
    check(False, "should raise RuntimeError for unknown extension")
    _any_fail = True
except RuntimeError:
    check(True, "RuntimeError raised for unknown extension")
except Exception as e:
    check(False, f"unexpected exception type: {type(e).__name__}: {e}")
    _any_fail = True


# --------------------------------------------------------------
# T16. FileConverter.html_to_text
# --------------------------------------------------------------

section("T16. FileConverter.html_to_text")
print("  Expected: tags stripped, readable text returned")
print("  Requires: pip install beautifulsoup4 lxml")

if not _lib("bs4"):
    print("  SKIP -- beautifulsoup4 not installed")
else:
    sub("html_to_text -- tags stripped")
    try:
        text = _fc.html_to_text(fp("sample.html"))
        check(isinstance(text, str),            "returns str")
        check(bool(text.strip()),               "non-empty")
        check("<html" not in text.lower(),      "no <html> tag in output")
        check("<script" not in text.lower(),    "no <script> tag in output")
        print(f"    Preview: {text[:120].strip()!r}")
    except Exception as e:
        print(f"  EXCEPTION: {e}")
        _any_fail = True


# --------------------------------------------------------------
# T17. FileConverter.audio_to_transcript
# --------------------------------------------------------------

section("T17. FileConverter.audio_to_transcript")
print("  Expected: non-empty transcript for sample.mp3 and sample.wav")
print("  Requires: pip install faster-whisper  (downloads model on first run ~tens MB)")

if not _lib("faster_whisper"):
    print("  SKIP -- faster-whisper not installed")
else:
    for fname, ext in [("sample.mp3", "mp3"), ("sample.wav", "wav")]:
        sub(f"audio_to_transcript -- {fname}")
        try:
            text = _fc.audio_to_transcript(fp(fname))
            check(isinstance(text, str),    "returns str")
            check(bool(text.strip()),       "non-empty")
            print(f"    Transcript: {text[:120].strip()!r}")
        except Exception as e:
            print(f"    EXCEPTION: {e}")
            _any_fail = True

    sub("audio_to_transcript -- language param accepted without error")
    try:
        text = _fc.audio_to_transcript(fp("sample.mp3"), language="ru")
        check(isinstance(text, str), "returns str with explicit language")
    except Exception as e:
        print(f"    EXCEPTION: {e}")
        _any_fail = True


# --------------------------------------------------------------
# T18. FileConverter.export_pdf
# --------------------------------------------------------------

section("T18. FileConverter.export_pdf")
print("  Expected: PDF file created, non-trivial size, path returned")
print("  Requires: pip install xhtml2pdf")

if not _lib("xhtml2pdf"):
    print("  SKIP -- xhtml2pdf not installed")
else:
    sub("export_pdf -- plain text input")
    try:
        out = _fc.export_pdf("Hello, PDF world!\n\nSecond paragraph.")
        p = Path(out)
        check(p.exists(),           "output file exists")
        check(p.suffix == ".pdf",   "extension is .pdf")
        check(p.stat().st_size > 500, f"non-trivial size ({p.stat().st_size} bytes)")
        p.unlink(missing_ok=True)
    except Exception as e:
        print(f"    EXCEPTION: {e}")
        _any_fail = True

    sub("export_pdf -- markdown input (headers, bold)")
    try:
        md = "# Title\n\nSome **bold** text.\n\n## Section\n\n- item 1\n- item 2"
        out = _fc.export_pdf(md, title="Test Doc")
        p = Path(out)
        check(p.exists(),           "output file exists")
        check(p.stat().st_size > 500, "non-trivial size")
        p.unlink(missing_ok=True)
    except Exception as e:
        print(f"    EXCEPTION: {e}")
        _any_fail = True

    sub("export_pdf -- explicit output_path")
    try:
        with _tempfile.TemporaryDirectory() as tmp:
            out_path = str(Path(tmp) / "explicit.pdf")
            out = _fc.export_pdf("Test content.", output_path=out_path)
            check(out == out_path,              "returned path matches requested path")
            check(Path(out_path).exists(),      "file exists at requested path")
    except Exception as e:
        print(f"    EXCEPTION: {e}")
        _any_fail = True


# --------------------------------------------------------------
# T19. FileConverter.export_docx
# --------------------------------------------------------------

section("T19. FileConverter.export_docx")
print("  Expected: DOCX file created, headings parsed, path returned")
print("  Requires: pip install python-docx")

if not _lib("docx"):
    print("  SKIP -- python-docx not installed")
else:
    sub("export_docx -- plain text")
    try:
        out = _fc.export_docx("First paragraph.\n\nSecond paragraph.")
        p = Path(out)
        check(p.exists(),           "output file exists")
        check(p.suffix == ".docx",  "extension is .docx")
        check(p.stat().st_size > 1000, f"non-trivial size ({p.stat().st_size} bytes)")
        p.unlink(missing_ok=True)
    except Exception as e:
        print(f"    EXCEPTION: {e}")
        _any_fail = True

    sub("export_docx -- headings parsed (# ## ###)")
    try:
        import docx as _dx
        content = "# H1 heading\n## H2 heading\n### H3 heading\nPlain paragraph."
        out = _fc.export_docx(content, title="Doc Title")
        doc = _dx.Document(out)
        styles = [p.style.name for p in doc.paragraphs]
        check(any("Heading 1" in s for s in styles),
              f"Heading 1 style present (styles: {styles})")
        check(any("Heading 2" in s for s in styles),
              f"Heading 2 style present (styles: {styles})")
        check(any("Heading 3" in s for s in styles),
              f"Heading 3 style present (styles: {styles})")
        texts = [p.text for p in doc.paragraphs]
        check("H1 heading" in texts,   f"H1 text correct (texts: {texts})")
        check("H2 heading" in texts,   f"H2 text correct (texts: {texts})")
        check("H3 heading" in texts,   f"H3 text correct (texts: {texts})")
        check("Plain paragraph." in texts, f"plain paragraph present (texts: {texts})")
        check("Doc Title" in texts,    f"title added as heading (texts: {texts})")
        Path(out).unlink(missing_ok=True)
    except Exception as e:
        print(f"    EXCEPTION: {e}")
        _any_fail = True

    sub("export_docx -- explicit output_path")
    try:
        with _tempfile.TemporaryDirectory() as tmp:
            out_path = str(Path(tmp) / "explicit.docx")
            out = _fc.export_docx("Content.", output_path=out_path)
            check(out == out_path,          "returned path matches requested path")
            check(Path(out_path).exists(),  "file exists at requested path")
    except Exception as e:
        print(f"    EXCEPTION: {e}")
        _any_fail = True


# --------------------------------------------------------------
# T20. FileConverter.export_xlsx
# --------------------------------------------------------------

section("T20. FileConverter.export_xlsx")
print("  Expected: XLSX file created, data readable back, path returned")
print("  Requires: pip install openpyxl")

if not _lib("openpyxl"):
    print("  SKIP -- openpyxl not installed")
else:
    import openpyxl as _oxl

    sub("export_xlsx -- list of lists")
    try:
        rows = [["Alice", 30], ["Bob", 25], ["Carol", 35]]
        out = _fc.export_xlsx(rows, headers=["Name", "Age"])
        p = Path(out)
        check(p.exists(),          "output file exists")
        check(p.suffix == ".xlsx", "extension is .xlsx")
        wb = _oxl.load_workbook(str(p))
        ws = wb.active
        all_rows = list(ws.iter_rows(values_only=True))
        check(all_rows[0] == ("Name", "Age"),
              f"header row correct (got: {all_rows[0]!r})")
        check(all_rows[1][0] == "Alice",
              f"first data row, col 0: Alice (got: {all_rows[1][0]!r})")
        # export_xlsx намеренно сохраняет все значения как строки (str(c) в строке записи),
        # поэтому openpyxl возвращает "30" (str), а не 30 (int).
        _got_val = all_rows[1][1]
        check(_got_val == "30",
              f"first data row, col 1: '30' as str (got: {_got_val!r}, type: {type(_got_val).__name__})")
        check(len(all_rows) == 4,                  "4 rows total (1 header + 3 data)")
        wb.close()
        p.unlink(missing_ok=True)
    except Exception as e:
        print(f"    EXCEPTION: {e}")
        _any_fail = True

    sub("export_xlsx -- list of dicts (auto headers from keys)")
    try:
        rows = [{"city": "Moscow", "pop": 12_000_000}, {"city": "SPb", "pop": 5_000_000}]
        out = _fc.export_xlsx(rows)
        wb = _oxl.load_workbook(out)
        ws = wb.active
        all_rows = list(ws.iter_rows(values_only=True))
        check(all_rows[0] == ("city", "pop"),
              f"headers auto-detected from dict keys (got: {all_rows[0]!r})")
        check(all_rows[1][0] == "Moscow",
              f"first data row, city: Moscow (got: {all_rows[1][0]!r})")
        wb.close()
        Path(out).unlink(missing_ok=True)
    except Exception as e:
        print(f"    EXCEPTION: {e}")
        _any_fail = True

    sub("export_xlsx -- custom sheet name")
    try:
        out = _fc.export_xlsx([["x", "y"]], sheet_name="MySheet")
        wb = _oxl.load_workbook(out)
        check(wb.active.title == "MySheet",    "sheet name set correctly")
        wb.close()
        Path(out).unlink(missing_ok=True)
    except Exception as e:
        print(f"    EXCEPTION: {e}")
        _any_fail = True

    sub("export_xlsx -- explicit output_path")
    try:
        with _tempfile.TemporaryDirectory() as tmp:
            out_path = str(Path(tmp) / "explicit.xlsx")
            out = _fc.export_xlsx([["a", "b"]], output_path=out_path)
            check(out == out_path,          "returned path matches requested path")
            check(Path(out_path).exists(),  "file exists at requested path")
    except Exception as e:
        print(f"    EXCEPTION: {e}")
        _any_fail = True


# --------------------------------------------------------------
# Summary
# --------------------------------------------------------------

if _t_section is not None:
    print(f"  Time: {time.perf_counter() - _t_section:.3f} sec")
    _t_section = None

print(f"\n{'='*60}")
if _any_fail:
    print("  RESULT: some checks FAILED - see FAIL lines above")
else:
    print("  RESULT: all checks passed")
print(f"  Total time: {time.perf_counter() - _t_total:.3f} sec")
print("=" * 60)